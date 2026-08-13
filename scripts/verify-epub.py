#!/usr/bin/env python3
"""
Verifica el EPUB contra el manuscrito del que salió.

No hay epubcheck ni calibre en esta máquina, así que las comprobaciones van
explícitas. Lo que se revisa, en orden de gravedad:

  1. Integridad de contenido: cada párrafo del .docx tiene que aparecer en el
     EPUB. Es la comprobación que importa — un conversor que pierde el Capítulo
     9 en silencio es peor que uno que falla ruidosamente.
  2. Las 4 tablas, celda por celda.
  3. Las cifras que sostienen el método (prueba del mes 12, caso Ramírez,
     caso Laura). Si una se perdió o se transformó, aquí se ve.
  4. XHTML bien formado en los 44 archivos.
  5. Estructura EPUB: mimetype primero y sin comprimir, container.xml, y que
     todo lo que declara el OPF exista de verdad.

Uso: python3 scripts/verify-epub.py
"""
from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

import docx
from lxml import etree

RAIZ = Path(__file__).resolve().parent.parent
DOCX = RAIZ / "src/book/Deuda Fuera Paz Dentro v3.7.docx"
EPUB = RAIZ / "build/epub/deuda-fuera-paz-dentro.epub"
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

fallos: list[str] = []
avisos: list[str] = []


def check(ok: bool, etiqueta: str, detalle: str = "") -> bool:
    print(f"  [{'ok ' if ok else 'FALLA'}] {etiqueta}{(' — ' + detalle) if detalle else ''}")
    if not ok:
        fallos.append(etiqueta)
    return ok


def normaliza(t: str) -> str:
    """Para leer: espacios colapsados."""
    return re.sub(r"\s+", " ", t.replace("\xa0", " ")).strip()


def sin_espacios(t: str) -> str:
    """Para comparar.

    Word guarda un párrafo como varios runs (uno por trozo en negrita) y mete
    saltos de línea dentro; el EPUB los vuelve <br/> y nodos separados. Los dos
    dicen lo mismo pero con espacios distintos, así que comparar respetándolos
    da falsos positivos. Lo que hay que probar es que el texto está, no cómo
    quedó espaciado.
    """
    return re.sub(r"\s+", "", t.replace("\xa0", " "))


def main() -> int:
    if not EPUB.exists():
        print(f"No existe {EPUB}. Corre antes scripts/build-epub.py", file=sys.stderr)
        return 1

    z = zipfile.ZipFile(EPUB)
    nombres = z.namelist()

    # ── 5. Estructura del contenedor ──────────────────────────────────
    print("\n== estructura EPUB ==")
    info0 = z.infolist()[0]
    check(info0.filename == "mimetype", "el mimetype es el primer archivo", info0.filename)
    check(info0.compress_type == zipfile.ZIP_STORED, "el mimetype va sin comprimir")
    check(
        z.read("mimetype").decode() == "application/epub+zip",
        "el mimetype dice application/epub+zip",
    )
    check("META-INF/container.xml" in nombres, "existe META-INF/container.xml")

    opf_xml = etree.fromstring(z.read("OEBPS/content.opf"))
    ns = {"opf": "http://www.idpf.org/2007/opf"}
    hrefs = [i.get("href") for i in opf_xml.findall(".//opf:manifest/opf:item", ns)]
    faltan = [h for h in hrefs if f"OEBPS/{h}" not in nombres]
    check(not faltan, f"los {len(hrefs)} recursos del manifest existen", str(faltan[:3]))

    idrefs = [s.get("idref") for s in opf_xml.findall(".//opf:spine/opf:itemref", ns)]
    ids = {i.get("id") for i in opf_xml.findall(".//opf:manifest/opf:item", ns)}
    check(all(i in ids for i in idrefs), f"el spine ({len(idrefs)}) apunta a ids reales")
    check(
        opf_xml.find(".//opf:manifest/opf:item[@properties='nav']", ns) is not None,
        "hay documento de navegación (nav.xhtml)",
    )

    # ── 4. XHTML bien formado ─────────────────────────────────────────
    print("\n== XHTML ==")
    xhtmls = [n for n in nombres if n.endswith(".xhtml")]
    malos = []
    for n in xhtmls:
        try:
            etree.fromstring(z.read(n))
        except etree.XMLSyntaxError as e:
            malos.append(f"{n}: {e}")
    check(not malos, f"los {len(xhtmls)} XHTML están bien formados", str(malos[:2]))

    # Texto plano de todo el EPUB, para las comprobaciones de contenido
    texto_epub = ""
    for n in xhtmls:
        arbol = etree.fromstring(z.read(n))
        # sin separador: dentro de un párrafo los nodos son continuos
        texto_epub += "\n" + "".join(arbol.itertext())
    texto_epub_cmp = sin_espacios(texto_epub)

    # ── 1. Integridad de contenido ────────────────────────────────────
    print("\n== contenido ==")
    doc = docx.Document(str(DOCX))
    P = doc.paragraphs
    toc_ini = next(i for i, p in enumerate(P) if p.text.strip() == "Índice")
    toc_fin = next(i for i in range(toc_ini + 1, len(P)) if P[i].style.name == "Heading 1")

    perdidos = []
    revisados = 0
    for i, p in enumerate(P):
        if toc_ini <= i < toc_fin:
            continue
        t = normaliza(p.text)
        if len(t) < 25:  # los muy cortos generan falsos positivos por puntuación
            continue
        revisados += 1
        if sin_espacios(t) not in texto_epub_cmp:
            perdidos.append((i, t[:70]))
    check(not perdidos, f"los {revisados} párrafos largos del .docx están en el EPUB",
          f"{len(perdidos)} perdidos: {perdidos[:3]}")

    # ── 2. Tablas ─────────────────────────────────────────────────────
    print("\n== tablas ==")
    celdas_total, celdas_perdidas = 0, []
    for ti, tbl in enumerate(doc.tables):
        for fila in tbl.rows:
            for celda in fila.cells:
                t = normaliza(celda.text)
                if len(t) < 3:
                    continue
                celdas_total += 1
                if sin_espacios(t) not in texto_epub_cmp:
                    celdas_perdidas.append(f"tabla {ti + 1}: {t[:40]}")
    check(not celdas_perdidas, f"las {celdas_total} celdas de las 4 tablas están",
          str(celdas_perdidas[:3]))
    tablas_en_epub = sum(
        len(etree.fromstring(z.read(n)).findall(".//{http://www.w3.org/1999/xhtml}table"))
        for n in xhtmls
    )
    check(tablas_en_epub == 4, "hay 4 <table> reales, no párrafos aplanados",
          f"encontradas: {tablas_en_epub}")

    # ── 3. Las cifras que sostienen el método ─────────────────────────
    print("\n== cifras del método ==")
    cifras = {
        "prueba del mes 12 — ROI": "sus mínimos en el mes 12 son $331",
        "prueba del mes 12 — saldo menor": "sus mínimos son $540",
        "prueba del mes 12 — avalancha": "sus mínimos son $797",
        "Ramírez — flujo liberado": "Liberaron $695 al mes",
        "Ramírez — ROI de OneMain": "ROI 56.0%",
        "Laura — el reparto del pago": "$179 se los comía el interés",
        "las cuatro fases": "Déficit, Oxígeno, Bola de Nieve o Avalancha",
        "fórmula del ROI": "pago mensual × 12",
    }
    for etiqueta, aguja in cifras.items():
        check(sin_espacios(aguja) in texto_epub_cmp, etiqueta)

    # ── Extras informativos ───────────────────────────────────────────
    print("\n== otros ==")
    nav = etree.fromstring(z.read("OEBPS/nav.xhtml"))
    entradas = nav.findall(".//{http://www.w3.org/1999/xhtml}li")
    check(len(entradas) == len(xhtmls) - 1, f"el índice tiene {len(entradas)} entradas, "
          f"una por archivo de contenido")
    if "OEBPS/toc.ncx" in nombres:
        print("  [ok ] incluye toc.ncx (compatibilidad con lectores viejos)")
    if not re.search(r"ISBN|isbn", z.read("OEBPS/content.opf").decode()):
        avisos.append("el OPF usa un UUID como identificador: falta el ISBN definitivo")

    print("\n" + "─" * 60)
    for a in avisos:
        print(f"  AVISO: {a}")
    if fallos:
        print(f"  {len(fallos)} comprobaciones FALLARON: {fallos}")
        return 1
    print("  todas las comprobaciones pasaron")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
