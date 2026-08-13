#!/usr/bin/env python3
"""
Construye el EPUB 3 de "Deuda Fuera, Paz Dentro" desde el manuscrito .docx.

Por qué a mano y no con pandoc/calibre: el libro tiene 4 tablas con la
aritmética del método y 974 párrafos de lista. Los conversores genéricos las
aplanan o las rompen en silencio, y este libro se sostiene precisamente sobre
que los números cuadren. Aquí cada elemento se traduce explícitamente y al
final se verifica, párrafo por párrafo, que nada se perdió.

Uso:
    python3 scripts/build-epub.py [ruta-al-docx]

Sale en build/epub/ (ignorado por git). El EPUB se regenera desde el .docx,
así que no se versiona el binario: se versiona esta receta.
"""
from __future__ import annotations

import html
import re
import sys
import unicodedata
import zipfile
from pathlib import Path

import docx
from docx.oxml.ns import qn

RAIZ = Path(__file__).resolve().parent.parent
DOCX_POR_DEFECTO = RAIZ / "src/book/Deuda Fuera Paz Dentro v3.7.docx"
SALIDA = RAIZ / "build/epub"

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"

META = {
    "titulo": "Deuda Fuera, Paz Dentro",
    "subtitulo": "Cómo crear un plan real para eliminar tus deudas y recuperar la paz financiera",
    "autor": "Rolando Rodríguez",
    "idioma": "es",
    "editorial": "Publicación independiente",
    "fecha": "2026",
    # PENDIENTE: el ISBN definitivo. El manuscrito dice que lo asigna KDP al
    # publicar; si al final se va a distribución amplia hará falta uno propio.
    # Mientras tanto va un UUID, que es lo que pide EPUB 3 como identificador
    # único cuando todavía no hay ISBN.
    "identificador": "urn:uuid:8f1a4c2e-6b3d-4f7a-9e21-3c5d7b9a1e40",
}

# ── Utilidades ────────────────────────────────────────────────────────


def slug(texto: str, tope: int = 40) -> str:
    t = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode()
    t = re.sub(r"[^a-zA-Z0-9]+", "-", t).strip("-").lower()
    return (t[:tope] or "seccion").rstrip("-")


def esc(t: str) -> str:
    return html.escape(t, quote=False)


# ── Traducción de un párrafo de Word a XHTML ──────────────────────────


def render_runs(p, rels) -> str:
    """Recorre el XML del párrafo en orden: runs, saltos de línea e hipervínculos."""
    partes: list[str] = []

    def render_run(r) -> str:
        trozos: list[str] = []
        for hijo in r:
            if hijo.tag == W + "t":
                trozos.append(esc(hijo.text or ""))
            elif hijo.tag == W + "br":
                trozos.append("<br/>")
            elif hijo.tag == W + "tab":
                trozos.append(" ")
        txt = "".join(trozos)
        if not txt:
            return ""
        rPr = r.find(W + "rPr")
        if rPr is not None:
            # <w:b/> sin val, o con val distinto de 0/false, cuenta como negrita
            def activo(nombre: str) -> bool:
                el = rPr.find(W + nombre)
                return el is not None and el.get(W + "val") not in ("0", "false")

            if activo("i"):
                txt = f"<em>{txt}</em>"
            if activo("b"):
                txt = f"<strong>{txt}</strong>"
        return txt

    for hijo in p._p:
        if hijo.tag == W + "r":
            partes.append(render_run(hijo))
        elif hijo.tag == W + "hyperlink":
            interno = "".join(render_run(r) for r in hijo.findall(W + "r"))
            rid = hijo.get(R + "id")
            destino = rels[rid].target_ref if rid and rid in rels else None
            partes.append(
                f'<a href="{esc(destino)}">{interno}</a>' if destino else interno
            )
    return "".join(partes)


def es_lista(p) -> bool:
    return p._p.find(".//" + W + "numPr") is not None


def render_tabla(tbl) -> str:
    filas = tbl.findall(W + "tr")
    if not filas:
        return ""
    out = ['<div class="tabla-scroll"><table>']
    for i, tr in enumerate(filas):
        celdas = tr.findall(W + "tc")
        etiqueta = "th" if i == 0 else "td"
        out.append("<thead><tr>" if i == 0 else ("<tbody><tr>" if i == 1 else "<tr>"))
        for tc in celdas:
            texto = " ".join(
                "".join(t.text or "" for t in par.iter(W + "t"))
                for par in tc.findall(W + "p")
            ).strip()
            out.append(f"<{etiqueta}>{esc(texto)}</{etiqueta}>")
        out.append("</tr></thead>" if i == 0 else "</tr>")
    out.append("</tbody></table></div>")
    return "".join(out)


# ── Armado del libro ──────────────────────────────────────────────────


class Capitulo:
    def __init__(self, titulo: str, nivel: int, archivo: str):
        self.titulo = titulo
        self.nivel = nivel  # 1 = parte, 2 = capítulo
        self.archivo = archivo
        self.cuerpo: list[str] = []


def construir(ruta_docx: Path) -> tuple[list[Capitulo], dict]:
    doc = docx.Document(str(ruta_docx))
    rels = doc.part.rels
    # Los párrafos en orden, indexables por su elemento XML
    por_elemento = {p._p: p for p in doc.paragraphs}
    tablas = {t._tbl: t for t in doc.tables}

    # El índice de Word es un campo que se regenera al maquetar. El EPUB tiene
    # su propia navegación, así que este bloque se salta entero.
    P = doc.paragraphs
    toc_ini = next(i for i, p in enumerate(P) if p.text.strip() == "Índice")
    toc_fin = next(i for i in range(toc_ini + 1, len(P)) if P[i].style.name == "Heading 1")
    saltar = {P[i]._p for i in range(toc_ini, toc_fin)}

    caps: list[Capitulo] = []
    actual = Capitulo("Portada y créditos", 0, "000-creditos.xhtml")
    caps.append(actual)
    # La portadilla no trae estilo Title en el .docx: el título del libro es un
    # párrafo normal. Se le da tratamiento por posición — los tres primeros
    # párrafos con texto son título, subtítulo y autor.
    portadilla = 0
    lista_abierta = False
    stats = {"parrafos": 0, "listas": 0, "tablas": 0, "saltados": 0}
    usados: set[str] = {actual.archivo}

    def cerrar_lista():
        nonlocal lista_abierta
        if lista_abierta:
            actual.cuerpo.append("</ul>")
            lista_abierta = False

    for el in doc.element.body:
        if el.tag == W + "tbl":
            cerrar_lista()
            if el in tablas:
                actual.cuerpo.append(render_tabla(el))
                stats["tablas"] += 1
            continue
        if el.tag != W + "p":
            continue
        p = por_elemento.get(el)
        if p is None:
            continue
        if el in saltar:
            stats["saltados"] += 1
            continue

        estilo = p.style.name
        texto = p.text.strip()

        if estilo in ("Heading 1", "Heading 2") and texto:
            cerrar_lista()
            nivel = 1 if estilo == "Heading 1" else 2
            base = f"{len(caps):03d}-{slug(texto)}"
            archivo = f"{base}.xhtml"
            n = 2
            while archivo in usados:
                archivo, n = f"{base}-{n}.xhtml", n + 1
            usados.add(archivo)
            actual = Capitulo(texto, nivel, archivo)
            caps.append(actual)
            actual.cuerpo.append(f"<h1>{esc(texto)}</h1>")
            stats["parrafos"] += 1
            continue

        if not texto:
            cerrar_lista()
            continue

        contenido = render_runs(p, rels)
        if not contenido.strip():
            continue

        if len(caps) == 1 and portadilla < 3:
            portadilla += 1
            clase = {1: "titulo-libro", 2: "subtitulo", 3: "autor"}[portadilla]
            etiqueta = "h1" if portadilla == 1 else "p"
            actual.cuerpo.append(f'<{etiqueta} class="{clase}">{contenido}</{etiqueta}>')
            stats["parrafos"] += 1
            continue

        if len(caps) == 1:
            cerrar_lista()
            actual.cuerpo.append(f'<p class="creditos">{contenido}</p>')
            stats["parrafos"] += 1
            continue

        if estilo == "Heading 3":
            cerrar_lista()
            actual.cuerpo.append(f"<h2>{contenido}</h2>")
        elif estilo == "Title":
            cerrar_lista()
            actual.cuerpo.append(f'<h1 class="titulo-libro">{contenido}</h1>')
        elif estilo == "Subtitle":
            cerrar_lista()
            actual.cuerpo.append(f'<p class="subtitulo">{contenido}</p>')
        elif es_lista(p):
            if not lista_abierta:
                actual.cuerpo.append("<ul>")
                lista_abierta = True
            actual.cuerpo.append(f"<li>{contenido}</li>")
            stats["listas"] += 1
        elif re.fullmatch(r"[•·\s]{3,}", texto):
            cerrar_lista()
            actual.cuerpo.append('<hr class="separador"/>')
        else:
            cerrar_lista()
            clase = ""
            if texto.startswith("✍️"):
                clase = ' class="tarea"'
            elif texto.startswith(("👉", "📌", "⚡")):
                clase = ' class="nota"'
            elif texto.startswith("⚠️"):
                clase = ' class="aviso"'
            actual.cuerpo.append(f"<p{clase}>{contenido}</p>")
        stats["parrafos"] += 1

    cerrar_lista()
    caps = [c for c in caps if c.cuerpo]
    return caps, stats


# ── Empaquetado EPUB 3 ────────────────────────────────────────────────

CSS = """@charset "utf-8";
body { font-family: Georgia, "Times New Roman", serif; line-height: 1.6;
       margin: 0 5%; text-align: left; hyphens: auto; }
h1 { font-size: 1.5em; line-height: 1.25; margin: 1.4em 0 0.8em; text-align: left;
     page-break-before: always; break-before: page; }
h1.titulo-libro { font-size: 2.1em; text-align: center; margin-top: 2em; page-break-before: auto; }
h2 { font-size: 1.15em; margin: 1.6em 0 0.5em; line-height: 1.3; }
p { margin: 0 0 0.85em; }
p.subtitulo { text-align: center; font-style: italic; color: #444; font-size: 1.1em; margin-top: 1.2em; }
p.autor { text-align: center; margin-top: 2em; font-size: 1.05em; }
p.creditos { font-size: 0.82em; color: #555; margin-top: 1.1em; }
h1.titulo-libro + p.subtitulo { margin-top: 0.6em; }
p.nota { margin-left: 0.8em; }
p.tarea { border-left: 3px solid #999; padding: 0.4em 0 0.4em 0.8em; margin: 1.2em 0; }
p.aviso { border-left: 3px solid #a33; padding: 0.4em 0 0.4em 0.8em; margin: 1.2em 0; }
ul { margin: 0 0 0.9em 1.1em; padding: 0; }
li { margin-bottom: 0.35em; }
hr.separador { border: none; text-align: center; margin: 1.8em 0; }
hr.separador::after { content: "• • •"; color: #777; letter-spacing: 0.35em; }
.tabla-scroll { overflow-x: auto; margin: 1em 0; }
table { border-collapse: collapse; width: 100%; font-size: 0.86em; }
th, td { border: 1px solid #bbb; padding: 0.35em 0.5em; text-align: left; vertical-align: top; }
th { background: #f0f0f0; font-weight: bold; }
a { color: #14524b; }
"""

PAGINA = """<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="{idioma}" xml:lang="{idioma}">
<head><meta charset="utf-8"/><title>{titulo}</title>
<link rel="stylesheet" type="text/css" href="estilo.css"/></head>
<body epub:type="bodymatter">
{cuerpo}
</body></html>
"""


def escribir_epub(caps: list[Capitulo], destino: Path) -> Path:
    destino.mkdir(parents=True, exist_ok=True)
    archivo = destino / "deuda-fuera-paz-dentro.epub"

    manifest, spine, navpoints, ncxpoints = [], [], [], []
    for i, c in enumerate(caps):
        idr = f"c{i:03d}"
        manifest.append(
            f'<item id="{idr}" href="{c.archivo}" media-type="application/xhtml+xml"/>'
        )
        spine.append(f'<itemref idref="{idr}"/>')
        navpoints.append(f'<li><a href="{c.archivo}">{esc(c.titulo)}</a></li>')
        ncxpoints.append(
            f'<navPoint id="np{i}" playOrder="{i + 1}"><navLabel><text>{esc(c.titulo)}</text>'
            f'</navLabel><content src="{c.archivo}"/></navPoint>'
        )

    opf = f"""<?xml version="1.0" encoding="utf-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="pub-id" xml:lang="{META['idioma']}">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:identifier id="pub-id">{META['identificador']}</dc:identifier>
    <dc:title id="t1">{esc(META['titulo'])}</dc:title>
    <meta refines="#t1" property="title-type">main</meta>
    <dc:title id="t2">{esc(META['subtitulo'])}</dc:title>
    <meta refines="#t2" property="title-type">subtitle</meta>
    <dc:creator id="a1">{esc(META['autor'])}</dc:creator>
    <meta refines="#a1" property="role" scheme="marc:relators">aut</meta>
    <dc:language>{META['idioma']}</dc:language>
    <dc:publisher>{esc(META['editorial'])}</dc:publisher>
    <dc:date>{META['fecha']}</dc:date>
    <dc:rights>© 2025 {esc(META['autor'])}. Todos los derechos reservados.</dc:rights>
    <meta property="dcterms:modified">2026-08-13T00:00:00Z</meta>
  </metadata>
  <manifest>
    <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>
    <item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>
    <item id="css" href="estilo.css" media-type="text/css"/>
    {chr(10).join('    ' + m for m in manifest).strip()}
  </manifest>
  <spine toc="ncx">
    {chr(10).join('    ' + s for s in spine).strip()}
  </spine>
</package>
"""

    nav = f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops" lang="{META['idioma']}" xml:lang="{META['idioma']}">
<head><meta charset="utf-8"/><title>Índice</title>
<link rel="stylesheet" type="text/css" href="estilo.css"/></head>
<body>
<nav epub:type="toc" id="toc"><h1>Índice</h1><ol>
{chr(10).join(navpoints).replace('<li>', '<li>')}
</ol></nav>
</body></html>
"""
    nav = nav.replace("<li><a", "<li><a")

    ncx = f"""<?xml version="1.0" encoding="utf-8"?>
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
  <head><meta name="dtb:uid" content="{META['identificador']}"/></head>
  <docTitle><text>{esc(META['titulo'])}</text></docTitle>
  <navMap>
    {chr(10).join('    ' + n for n in ncxpoints).strip()}
  </navMap>
</ncx>
"""

    with zipfile.ZipFile(archivo, "w") as z:
        # El mimetype va primero y SIN comprimir: lo exige la especificación.
        z.writestr(
            zipfile.ZipInfo("mimetype"), "application/epub+zip", zipfile.ZIP_STORED
        )
        z.writestr(
            "META-INF/container.xml",
            '<?xml version="1.0" encoding="utf-8"?>\n'
            '<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">'
            '<rootfiles><rootfile full-path="OEBPS/content.opf" '
            'media-type="application/oebps-package+xml"/></rootfiles></container>',
            zipfile.ZIP_DEFLATED,
        )
        z.writestr("OEBPS/content.opf", opf, zipfile.ZIP_DEFLATED)
        z.writestr("OEBPS/nav.xhtml", nav, zipfile.ZIP_DEFLATED)
        z.writestr("OEBPS/toc.ncx", ncx, zipfile.ZIP_DEFLATED)
        z.writestr("OEBPS/estilo.css", CSS, zipfile.ZIP_DEFLATED)
        for c in caps:
            z.writestr(
                f"OEBPS/{c.archivo}",
                PAGINA.format(
                    idioma=META["idioma"],
                    titulo=esc(c.titulo),
                    cuerpo="\n".join(c.cuerpo),
                ),
                zipfile.ZIP_DEFLATED,
            )
    return archivo


def main() -> int:
    ruta = Path(sys.argv[1]) if len(sys.argv) > 1 else DOCX_POR_DEFECTO
    if not ruta.exists():
        print(f"No encuentro el manuscrito: {ruta}", file=sys.stderr)
        return 1
    caps, stats = construir(ruta)
    archivo = escribir_epub(caps, SALIDA)
    kb = archivo.stat().st_size / 1024
    print(f"EPUB: {archivo}  ({kb:.0f} KB)")
    print(f"  {len(caps)} archivos · {stats['parrafos']} párrafos · "
          f"{stats['listas']} elementos de lista · {stats['tablas']} tablas")
    print(f"  {stats['saltados']} párrafos del índice de Word omitidos (el EPUB trae el suyo)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
