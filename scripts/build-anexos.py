#!/usr/bin/env python3
"""
Genera los tres anexos que acompañan al libro, desde el propio manuscrito.

Se extraen del .docx en vez de escribirse aparte a propósito: un anexo que se
redacta por su cuenta se desincroniza del libro a la primera revisión. Así, si
el Capítulo 9 cambia, la guía cambia con él.

Los tres:
  1. guia-estrategias.pdf   — las cuatro fases y las reglas de decisión (Cap. 9)
  2. scripts-negociacion.pdf — los guiones de llamada y la goodwill letter (Cap. 13)
  3. calendario-7-3-1.ics    — los tres recordatorios de pago

Sustituyen a los placeholders de public/downloads/ (946 y 482 bytes).

Uso: python3 scripts/build-anexos.py
"""
from __future__ import annotations

import importlib.util
import sys
import tempfile
from pathlib import Path

import docx

RAIZ = Path(__file__).resolve().parent.parent
SALIDA = RAIZ / "build/anexos"

_spec = importlib.util.spec_from_file_location("build_epub", Path(__file__).parent / "build-epub.py")
_be = importlib.util.module_from_spec(_spec)
assert _spec.loader
_spec.loader.exec_module(_be)

W = _be.W

CSS = """
@page { size: letter; margin: 0.8in 0.9in; }
body { font-family: Georgia, "Times New Roman", serif; font-size: 11pt; line-height: 1.5;
       color: #111; margin: 0; }
.cabecera { border-bottom: 2px solid #14524b; padding-bottom: 0.5em; margin-bottom: 1.4em; }
.cabecera .marca { font-family: system-ui, sans-serif; font-size: 8.5pt; letter-spacing: .13em;
                   text-transform: uppercase; color: #14524b; margin: 0 0 .35em; }
h1 { font-size: 19pt; margin: 0 0 .25em; line-height: 1.15; }
.cabecera p.sub { margin: 0; color: #555; font-size: 10pt; font-style: italic; }
h2 { font-size: 12.5pt; margin: 1.5em 0 .4em; page-break-after: avoid; }
p { margin: 0 0 .55em; }
ul { margin: 0 0 .7em 1.1em; padding: 0; }
li { margin-bottom: .25em; }
p.nota { margin-left: .5em; }
p.tarea, p.aviso { border-left: 2.5px solid #888; padding: .3em 0 .3em .7em; margin: .9em 0; }
p.aviso { border-left-color: #a33; }
table { border-collapse: collapse; width: 100%; font-size: 9pt; margin: .8em 0;
        page-break-inside: avoid; }
th, td { border: .5pt solid #999; padding: .3em .45em; text-align: left; vertical-align: top; }
th { background: #eef2f1; }
hr.separador { border: none; margin: 1.2em 0; }
hr.separador::after { content: "• • •"; color: #888; letter-spacing: .3em; }
.pie { margin-top: 2.2em; padding-top: .7em; border-top: 1px solid #ccc;
       font-family: system-ui, sans-serif; font-size: 8.5pt; color: #666; }
"""


def secciones(doc) -> list[tuple[int, str, int]]:
    """(índice, texto, nivel) de cada encabezado. Nivel 4 = párrafo normal."""
    out = []
    for i, p in enumerate(doc.paragraphs):
        n = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3}.get(p.style.name)
        if n:
            out.append((i, p.text.strip(), n))
    return out


def rango(doc, titulo: str, hasta: str | None = None) -> tuple[int, int]:
    """Del encabezado `titulo` hasta `hasta` (o hasta el siguiente de nivel ≤).

    Acepta también anclas que en el .docx quedaron como párrafo normal aunque
    se lean como subtítulo — "Las cuatro fases" es una de ellas. En ese caso el
    bloque va hasta el siguiente encabezado real.
    """
    heads = secciones(doc)
    ini = next((i for i, t, _ in heads if t == titulo), None)
    if ini is None:
        ini = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == titulo), None)
        if ini is None:
            raise SystemExit(f"No encuentro la sección {titulo!r} en el manuscrito")
        fin = next((i for i, _, _ in heads if i > ini), len(doc.paragraphs))
        return ini, fin
    nivel = next(n for i, t, n in heads if i == ini)
    if hasta:
        fin = next((i for i, t, _ in heads if t == hasta and i > ini), None)
        if fin is None:
            raise SystemExit(f"No encuentro la sección {hasta!r}")
        return ini, fin
    fin = next((i for i, t, n in heads if i > ini and n <= nivel), len(doc.paragraphs))
    return ini, fin


def render(doc, desde: int, hasta: int, saltar_titulo: bool = True) -> str:
    """Traduce un rango de párrafos a XHTML, reusando el motor del EPUB."""
    rels = doc.part.rels
    P = doc.paragraphs
    out: list[str] = []
    lista = False
    for i in range(desde + (1 if saltar_titulo else 0), hasta):
        p = P[i]
        texto = p.text.strip()
        if not texto:
            if lista:
                out.append("</ul>")
                lista = False
            continue
        contenido = _be.render_runs(p, rels)
        if not contenido.strip():
            continue
        if _be.es_lista(p):
            if not lista:
                out.append("<ul>")
                lista = True
            out.append(f"<li>{contenido}</li>")
            continue
        if lista:
            out.append("</ul>")
            lista = False
        if p.style.name in ("Heading 2", "Heading 3"):
            out.append(f"<h2>{contenido}</h2>")
        elif __import__("re").fullmatch(r"[•·\s]{3,}", texto):
            out.append('<hr class="separador"/>')
        else:
            clase = ""
            if texto.startswith("✍️"):
                clase = ' class="tarea"'
            elif texto.startswith(("👉", "📌", "⚡")):
                clase = ' class="nota"'
            elif texto.startswith("⚠️"):
                clase = ' class="aviso"'
            out.append(f"<p{clase}>{contenido}</p>")
    if lista:
        out.append("</ul>")
    return "\n".join(out)


def pagina(titulo: str, subtitulo: str, cuerpo: str) -> str:
    return f"""<!doctype html><html lang="es"><head><meta charset="utf-8">
<title>{_be.esc(titulo)}</title><style>{CSS}</style></head><body>
<header class="cabecera">
  <p class="marca">Deuda Fuera, Paz Dentro</p>
  <h1>{_be.esc(titulo)}</h1>
  <p class="sub">{_be.esc(subtitulo)}</p>
</header>
{cuerpo}
<p class="pie">Extracto de <em>Deuda Fuera, Paz Dentro</em>, de Rolando Rodríguez.
Calcula tu IPD gratis en deudafuerapazdentro.com/diagnostico · Material educativo:
no constituye asesoría financiera, legal ni fiscal.</p>
</body></html>"""


ICS = """BEGIN:VCALENDAR
VERSION:2.0
PRODID:-//Deuda Fuera Paz Dentro//Calendario 7-3-1//ES
CALSCALE:GREGORIAN
METHOD:PUBLISH
X-WR-CALNAME:Pagos 7-3-1
X-WR-TIMEZONE:America/Chicago
{eventos}END:VCALENDAR
"""

EVENTO = """BEGIN:VEVENT
UID:7-3-1-{n}@deudafuerapazdentro.com
DTSTAMP:20260101T090000Z
DTSTART;VALUE=DATE:{fecha}
DURATION:P1D
RRULE:FREQ=MONTHLY
SUMMARY:{titulo}
DESCRIPTION:{desc}
TRANSP:TRANSPARENT
BEGIN:VALARM
TRIGGER:-PT9H
ACTION:DISPLAY
DESCRIPTION:{titulo}
END:VALARM
END:VEVENT
"""


def plegar(ics: str) -> str:
    """Pliega a 75 octetos, como manda RFC 5545.

    Las líneas largas sin plegar hacen que algunos clientes trunquen la
    descripción o descarten el evento entero.
    """
    salida: list[str] = []
    for linea in ics.splitlines():
        b = linea.encode("utf-8")
        if len(b) <= 75:
            salida.append(linea)
            continue
        trozo, actual = [], b""
        for ch in linea:
            e = ch.encode("utf-8")
            # 74 en las continuaciones: el espacio inicial también cuenta
            if len(actual) + len(e) > (75 if not trozo else 74):
                trozo.append(actual.decode("utf-8"))
                actual = b""
            actual += e
        trozo.append(actual.decode("utf-8"))
        salida.append(trozo[0])
        salida.extend(" " + t for t in trozo[1:])
    return "\r\n".join(salida) + "\r\n"


def calendario() -> str:
    """El 7-3-1 del libro: revisar, ejecutar, verificar.

    El .ics anterior no traía DTSTART, y sin eso la mayoría de los calendarios
    rechaza el evento entero: el archivo se veía bien y no importaba nada.
    Las fechas son de referencia (día 1 de un mes tipo); quien lo importe las
    mueve a su fecha de corte.
    """
    pasos = [
        ("20260124", "Pago: revisar (7 días antes)",
         "Revisa saldos y fondos. Si algo no alcanza, hoy hay tiempo de moverlo."),
        ("20260128", "Pago: ejecutar (3 días antes)",
         "Programa o haz el pago. No lo dejes para el último día."),
        ("20260130", "Pago: verificar (1 día antes)",
         "Confirma que el pago se acreditó. Un pago programado que falla se ve igual que uno hecho."),
    ]
    ev = "".join(
        EVENTO.format(n=i + 1, fecha=f, titulo=t, desc=d)
        for i, (f, t, d) in enumerate(pasos)
    )
    return plegar(ICS.format(eventos=ev))


def a_pdf(htmls: list[tuple[str, Path]]) -> None:
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        nav = p.chromium.launch()
        pg = nav.new_page()
        for html, destino in htmls:
            with tempfile.NamedTemporaryFile("w", suffix=".html", delete=False, encoding="utf-8") as f:
                f.write(html)
                tmp = Path(f.name)
            try:
                pg.goto(tmp.as_uri())
                pg.wait_for_load_state("networkidle")
                pg.pdf(path=str(destino), format="Letter", print_background=True,
                       prefer_css_page_size=True, outline=True, tagged=True)
            finally:
                tmp.unlink(missing_ok=True)
        nav.close()


def main() -> int:
    ruta = Path(sys.argv[1]) if len(sys.argv) > 1 else _be.DOCX_POR_DEFECTO
    if not ruta.exists():
        print(f"No encuentro el manuscrito: {ruta}", file=sys.stderr)
        return 1
    doc = docx.Document(str(ruta))
    SALIDA.mkdir(parents=True, exist_ok=True)

    # 1. Guía de estrategias: las cuatro fases + las reglas de decisión.
    f_ini, f_fin = rango(doc, "Las cuatro fases")
    r_ini, r_fin = rango(doc, "Reglas de decisión (resumen operativo)")
    guia = pagina(
        "Guía de Estrategias",
        "Las cuatro fases del Selector: en cuál estás y con qué criterio pagar en cada una.",
        f"<h2>Las cuatro fases</h2>\n{render(doc, f_ini, f_fin)}\n"
        f"<h2>Reglas de decisión</h2>\n{render(doc, r_ini, r_fin)}",
    )

    # 2. Scripts de negociación: del Capítulo 13 completo, sin los cierres.
    s_ini, s_fin = rango(doc, "Capítulo 13: Negociación con acreedores",
                         "Capítulo 14: Cómo reparar tu historial y dejar de ser castigado por el pasado")
    scripts = pagina(
        "Scripts para negociar con acreedores",
        "Qué decir, en qué orden y cuándo aceptar el trato.",
        render(doc, s_ini, s_fin),
    )

    pdf_guia = SALIDA / "guia-estrategias.pdf"
    pdf_scripts = SALIDA / "scripts-negociacion.pdf"
    a_pdf([(guia, pdf_guia), (scripts, pdf_scripts)])

    # 3. Calendario.
    ics = SALIDA / "calendario-7-3-1.ics"
    ics.write_text(calendario(), encoding="utf-8")

    for f in (pdf_guia, pdf_scripts, ics):
        print(f"  {f.name:26s} {f.stat().st_size / 1024:6.1f} KB")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
