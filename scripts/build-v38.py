"""Genera el v3.8 desde el v3.7: los cinco arreglos de imprenta y la oferta de /libro.

Cada cambio verifica su ancla antes de tocar nada; si el v3.7 no es el que se
espera, el script aborta en vez de escribir a ciegas. Nunca sobreescribe el
original: siempre escribe un archivo de versión nueva.
"""
import sys
import docx

ORIGEN = "src/book/Deuda Fuera Paz Dentro v3.7.docx"
DESTINO = "src/book/Deuda Fuera Paz Dentro v3.8.docx"

d = docx.Document(ORIGEN)
ps = d.paragraphs
cambios = []


def ancla(i: int, fragmento: str, estilo: str | None = None) -> None:
    """Aborta si el párrafo no es el que este script cree estar editando."""
    if fragmento not in ps[i].text:
        sys.exit(f"ABORTA: ¶{i} no contiene {fragmento!r}. Es: {ps[i].text[:90]!r}")
    if estilo and ps[i].style.name != estilo:
        sys.exit(f"ABORTA: ¶{i} tiene estilo {ps[i].style.name!r}, se esperaba {estilo!r}")


# ── 1. El año de copyright ────────────────────────────────────────────────
# Primera edición: 2026, pero el aviso decía 2025 en dos sitios.
for i in (6, 2613):
    ancla(i, "© 2025 Rolando Rodríguez")
    ps[i].runs[0].text = ps[i].runs[0].text.replace("© 2025", "© 2026")
    cambios.append(f"¶{i}  © 2025 → © 2026")

# ── 2. Los metadatos del documento ────────────────────────────────────────
# Esto es lo que leen KDP y los conversores a EPUB, y decía «Word Document»
# sin autor ni idioma. Era la causa; el estilo del título era el síntoma.
cp = d.core_properties
cp.title = "Deuda Fuera, Paz Dentro"
cp.author = "Rolando Rodríguez"
cp.language = "es"
cp.subject = "Finanzas personales · eliminación de deudas"
cambios.append("meta  title 'Word Document' → 'Deuda Fuera, Paz Dentro'")
cambios.append("meta  author y language, que estaban vacíos")

# El título del libro iba como texto corriente. Se le da el estilo Title,
# que es a lo que recurre un conversor si los metadatos fallan.
# «Índice» se queda como Title a propósito: el índice es un campo TOC \o "1-2",
# así que un Heading 1 ahí lo metería dentro de su propio índice.
ancla(2, "Deuda Fuera", "normal")
ancla(59, "Índice", "Title")
ps[2].style = d.styles["Title"]
cambios.append("¶2   'Deuda Fuera, Paz Dentro'  normal → Title")

# ── 3. La línea de ISBN ───────────────────────────────────────────────────
# Era una frase de relleno que se habría impreso literal. KDP asigna el ISBN
# gratuito durante el alta del libro en papel, antes de subir el interior
# definitivo: entonces se pega el número real.
ancla(8, "ISBN: asignado por Amazon KDP")
t = ps[8].runs[0].text
ps[8].runs[0].text = "\n".join(l for l in t.split("\n") if not l.startswith("ISBN:"))
cambios.append("¶8   fuera la frase de relleno del ISBN")

# ── 4. El nivel de «No leas esto solo» ────────────────────────────────────
# Era Heading 2 entre hermanos Heading 1, así que en el índice generado
# quedaba anidado bajo «Sobre el autor».
ancla(2594, "No leas esto solo", "Heading 2")
ps[2594].style = d.styles["Heading 1"]
cambios.append("¶2594 'No leas esto solo'  Heading 2 → Heading 1")

# ── 5. La oferta del mes gratis ───────────────────────────────────────────
# Va en el material del final, muy pasado el umbral que Amazon enseña en
# «Echa un vistazo»: un enlace impreso en las primeras páginas sería público
# desde el primer día.
ancla(2546, "Es un libro para usar.")
ancla(2547, "Glosario", "Heading 1")

BLOQUE = [
    ("Heading 1", "Tu primer mes de la herramienta va incluido"),
    ("normal",
     "Todo lo que acabas de leer se calcula solo en el GPS Anti-Deuda: cargas tus "
     "deudas y te dice tu IPD, en qué fase estás y con qué criterio pagar. Es la "
     "misma aritmética de estos capítulos, aplicada a tus números, en quince minutos."),
    ("normal", "Tienes un mes completo, sin costo y sin tarjeta:"),
    ("normal", "deudafuerapazdentro.com/libro"),
    ("normal",
     "Entras con tu correo, cargas tus deudas y ves tu plan. Si al final del mes "
     "no es para ti, no haces nada y se acaba."),
]
glosario = ps[2547]
for estilo, texto in BLOQUE:
    glosario.insert_paragraph_before(texto, style=d.styles[estilo])
cambios.append(f"¶2547 insertados {len(BLOQUE)} párrafos: la oferta de /libro")

d.save(DESTINO)

print(f"Escrito {DESTINO}\n")
for c in cambios:
    print(f"  · {c}")

nuevo = docx.Document(DESTINO)
print(f"\n  párrafos: {len(ps)} → {len(nuevo.paragraphs)}")
print(f"  palabras: {sum(len(p.text.split()) for p in nuevo.paragraphs):,}")
print(f"  tablas:   {len(nuevo.tables)}")
